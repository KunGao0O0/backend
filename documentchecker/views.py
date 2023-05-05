from django.http import HttpResponse
from rest_framework import generics, status
from rest_framework.response import Response
from .models import Task, File
from .tasks import similaritycheck
from .serializers import FileSerializer, TaskSerialiazer
from django.shortcuts import render
from rest_framework.views import APIView
from .models import Threshold
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.http.request import QueryDict
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.oxml import parse_xml
from lxml import etree
from datetime import datetime
import json
import io
import os

import openai
import docx



# Create your views here.


def index(request):
    return render(request, "index.html")


############################################################################################
class DriveView(generics.ListAPIView, generics.GenericAPIView):
    '''
    This class receives http request about google drive
    '''
    authentication_classes = []
    permission_classes = []
    serializer_class = FileSerializer
    queryset = File.objects.all().order_by("-id")
    lookup_field = "id"

    def post(self, request):
        file = request.data.get("file", None)

        # Set up the OAuth2 authentication flow

        print("request: ", request.data)

        access_token = request.data.get('accessToken')
        auther = request.data.get('owners')[0].get('displayName')
        create_time = request.data.get('createdTime')
        modified_time = request.data.get('modifiedTime')
        modified_person = request.data.get('lastModifyingUser').get("displayName")
        creds = Credentials(token=access_token, scopes=['https://www.googleapis.com/auth/drive'])

        drive_service = build('drive', 'v3', credentials=creds)

        results = drive_service.files().list(
            fields="files(name,id,createdTime, modifiedTime, owners, lastModifyingUser)"
        ).execute()

        # Download a file from Google Drive
        file_id = request.data.get('fileid')
        request2 = drive_service.files().get_media(fileId=file_id)
        file_object = io.BytesIO()
        downloader = MediaIoBaseDownload(file_object, request2)
        done = False
        while done is False:
            stat, done = downloader.next_chunk()
            print("Download %d%%." % int(stat.progress() * 100))

        # Save the downloaded file to disk
        with open('tem.docx', 'wb') as f:
            f.write(file_object.getbuffer())

        def update_core_properties(docx_path, new_author, new_created_time, new_modified_time, last_modified_persion):

            """
            This function aims to change the metadata of docx but it seems that it doesn't work....
            """
            doc = Document(docx_path)
            core_properties = doc.core_properties
            core_properties.author = new_author
            core_properties.created = new_created_time
            core_properties.modified = new_modified_time
            core_properties.last_modified_by = last_modified_persion

            doc.save(docx_path)


            # Replace with the path to your .docx file

        input_docx_path = 'tem.docx'
        date_format = "%Y-%m-%dT%H:%M:%S.%fZ"
        # Update the core properties with the new author and created time

        new_author = auther
        new_created_time = datetime.strptime(create_time, date_format)
        new_modified_time = datetime.strptime(modified_time, date_format)
        last_modified_persion = modified_person

        update_core_properties(input_docx_path, new_author, new_created_time, new_modified_time, last_modified_persion)


        # Read the doc from disk and create a new InMemoryUploadedFile and pass it to serializer
        with open('tem.docx', 'rb') as f:
            file_content = f.read()

        file_object22 = io.BytesIO(file_content)

        uploaded_file = InMemoryUploadedFile(
            file_object22,
            field_name='file',
            name='tem.docx',
            content_type='application/msword',
            size=len(file_content),
            charset=None
        )

        query_dict = QueryDict(mutable=True)
        query_dict.update({'file': uploaded_file})
        temp = query_dict

        print("query_dict: ", query_dict)
        print("fileeee: ", query_dict.get("file"))

        if temp == "":
            return Response(
                {"file": "file is empty"}, status=status.HTTP_400_BAD_REQUEST
            )
        if temp is None:
            return Response(
                {"file": "file is none"}, status=status.HTTP_400_BAD_REQUEST
            )

        serializer = self.serializer_class(
            data=temp, context={"request": self.request}
        )
        print("is here")
        if serializer.is_valid(raise_exception=True):
            serializer.save()

            original_data = serializer.data
            # This part conduct gpi classification api invocation. Please update it to company api if possible.
            openai.api_key = "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"

            doc = docx.Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            file_text = '\n'.join(full_text)
            input_text = file_text

            if len(input_text) > 4000:
                input_text = input_text[:4000]

            attribute_dict = {
                "Personal": "Creates a sense of person/personality for the writer in the reader’s mind. Typically written in a very informal style with first and/or second person and emotive language.",
                "Imaginative": "Uses the writer’s creativity and imagination to entertain the reader. Uses techniques such as variation in sentence length, juxtaposition of different sentence lengths, careful control of structure and sequencing, to add to the overall effect by creating the desired atmosphere or conveying the required emotion.",
                "Persuasive": "Aims to convince readers about an idea, opinion or a course of action in order to achieve a desired outcome. Persuasive techniques chosen are influenced by the nature of the target audience; that is, the language (vocabulary, sentence structures, style/register), structure and sequencing of the piece are framed with the particular audience and purpose in mind",
                "Informative": "Generally uses facts, examples, explanations, analogies and sometimes statistical information, quotations and references as evidence. Chooses language, structure and sequence to make the message clear and unambiguous, so the sequencing of information is usually logical and predictable",
                "Evaluative": "Presents two or more important aspects of an issue or sides of an argument and discusses these rationally and objectively; using evidence to support the contrasting sides or alternatives. Uses objective style; appeals to reason not emotion; creation of an impression of balance and impartiality is essential.", }
            attribute_list = ["Geography, Agriculture and Environment", "Biology", "Mathematics",
                              "Culture, Arts and Music", "History", "Science",
                              "Education", "Health and Medicine", "English Language and Literature", "Technology",
                              "Physics", "Society", "Politics",
                              "Business and Economics", "Chemistry", "Legal studies", "Engineering", "Psychology",
                              "Architecture, Planning and Design",
                              "Journalism and Media", "Religious studies", "Other"]

            prompt1 = f"Which key in the dictionary {json.dumps(attribute_dict)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response1 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt1,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer1 = response1.choices[0].text.strip()
            category_key = None
            for option in attribute_dict.keys():
                if option.lower() in answer1.lower():
                    category_key = option
                    break
            prompt2 = f"Which category in the list {json.dumps(attribute_list)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response2 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt2,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer2 = response2.choices[0].text.strip()

            category_list = None
            for option in attribute_list:
                if option.lower() in answer2.lower():
                    category_list = option
                    break

            print(f"Category key: {category_key}")
            print(f"Category list: {category_list}")

            additional_data = {"category_key": category_key, "category_list": category_list}
            merged_data = {**original_data, **additional_data}
            return Response(merged_data, status=status.HTTP_201_CREATED)

        print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def get(self, request, *args, **kwargs):
        return self.retrieve(request, *args, **kwargs)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        response = HttpResponse(instance.file, content_type="text/txt")
        response["Content-Disposition"] = "attachment; filename={}".format(
            str(instance.file.name)
        )
        return response


class UploadFile(generics.ListAPIView, generics.GenericAPIView):
    authentication_classes = []
    permission_classes = []
    serializer_class = FileSerializer
    queryset = File.objects.all().order_by("-id")
    lookup_field = "id"



    def post(self, request):
        # print("file:", request)
        # print("---")

        print("file:", request.data)
        file = request.data.get("file", None)


        if file == "":
            return Response(
                {"file": "file is empty"}, status=status.HTTP_400_BAD_REQUEST
            )
        if file is None:
            return Response(
                {"file": "file is none"}, status=status.HTTP_400_BAD_REQUEST
            )

        serializer = self.serializer_class(
            data=request.data, context={"request": self.request}
        )
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            original_data = serializer.data
            
            # This part conduct gpi classification api invocation. Please update it to company api if possible.
            openai.api_key = "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"

            doc = docx.Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            file_text = '\n'.join(full_text)
            input_text = file_text

            if len(input_text) > 4000:
                input_text = input_text[:4000]

            attribute_dict = {
                "Personal": "Creates a sense of person/personality for the writer in the reader’s mind. Typically written in a very informal style with first and/or second person and emotive language.",
                "Imaginative": "Uses the writer’s creativity and imagination to entertain the reader. Uses techniques such as variation in sentence length, juxtaposition of different sentence lengths, careful control of structure and sequencing, to add to the overall effect by creating the desired atmosphere or conveying the required emotion.",
                "Persuasive": "Aims to convince readers about an idea, opinion or a course of action in order to achieve a desired outcome. Persuasive techniques chosen are influenced by the nature of the target audience; that is, the language (vocabulary, sentence structures, style/register), structure and sequencing of the piece are framed with the particular audience and purpose in mind",
                "Informative": "Generally uses facts, examples, explanations, analogies and sometimes statistical information, quotations and references as evidence. Chooses language, structure and sequence to make the message clear and unambiguous, so the sequencing of information is usually logical and predictable",
                "Evaluative": "Presents two or more important aspects of an issue or sides of an argument and discusses these rationally and objectively; using evidence to support the contrasting sides or alternatives. Uses objective style; appeals to reason not emotion; creation of an impression of balance and impartiality is essential.", }
            attribute_list = ["Geography, Agriculture and Environment", "Biology", "Mathematics",
                              "Culture, Arts and Music", "History", "Science",
                              "Education", "Health and Medicine", "English Language and Literature", "Technology",
                              "Physics", "Society", "Politics",
                              "Business and Economics", "Chemistry", "Legal studies", "Engineering", "Psychology",
                              "Architecture, Planning and Design",
                              "Journalism and Media", "Religious studies", "Other"]

            prompt1 = f"Which key in the dictionary {json.dumps(attribute_dict)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response1 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt1,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer1 = response1.choices[0].text.strip()
            category_key = None
            for option in attribute_dict.keys():
                if option.lower() in answer1.lower():
                    category_key = option
                    break
            prompt2 = f"Which category in the list {json.dumps(attribute_list)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response2 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt2,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer2 = response2.choices[0].text.strip()

            category_list = None
            for option in attribute_list:
                if option.lower() in answer2.lower():
                    category_list = option
                    break

            print(f"Category key: {category_key}")
            print(f"Category list: {category_list}")

            additional_data = {"category_key": category_key, "category_list": category_list}
            merged_data = {**original_data, **additional_data}
            return Response(merged_data, status=status.HTTP_201_CREATED)
        print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def get(self, request, *args, **kwargs):
        return self.retrieve(request, *args, **kwargs)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        response = HttpResponse(instance.file, content_type="text/txt")
        response["Content-Disposition"] = "attachment; filename={}".format(
            str(instance.file.name)
        )
        return response


class TaskView(
    generics.RetrieveAPIView, generics.UpdateAPIView, generics.GenericAPIView
):
    authentication_classes = []
    permission_classes = []
    serializer_class = TaskSerialiazer
    queryset = Task.objects.all().order_by("-id")
    lookup_field = "id"

    def post(self, request):
        files = request.data.get("file_id")
        print("post,here:", files)
        authors = request.data.get("authors")
        if len(files) < 2:
            return Response(
                {"file": "Add more files"}, status=status.HTTP_400_BAD_REQUEST
            )
        if files is None or files == []:
            return Response(
                {"file": "Select atleast two file"}, status=status.HTTP_400_BAD_REQUEST
            )
        if authors is None or len(authors) == 0:
            return Response(
                {"author": "Select atleast one author"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # authors = File.objects.filter(id__in=files).values_list('author', flat=True)
        # authors = list(set(authors))
        # if len(objts) <= 1:
        #     return Response(
        #         {"file": "Select atleast two files for {}".format(authors)},
        #         status=status.HTTP_400_BAD_REQUEST,
        #     )

        # create and task object
        SimilarityCheck_obj = Task()
        SimilarityCheck_obj.authors = json.dumps(authors)
        SimilarityCheck_obj.save()

        # celery task
        similaritycheck.delay(file=files, authors=authors, id=SimilarityCheck_obj.id)

        return Response({"file": "Files submitted sucsessfully", "task_id": SimilarityCheck_obj.id},
                        status=status.HTTP_200_OK)

    def get(self, request, *args, **kwargs):
        print("get,here:", request)
        return self.retrieve(request, *args, **kwargs)

    # def patch(self, request, *args, **kwargs):
    #     return self.partial_update(request, *args, **kwargs)

    # def update(self, request, *args, **kwargs):
    #     partial = kwargs.pop("partial", False)
    #     instance = self.get_object()
    #     completed_file = instance.completed_file
    #     threshold_file = instance.threshold_file
    #     if completed_file < threshold_file:
    #         return Response(
    #             {"file": "upload more file "}, status=status.HTTP_400_BAD_REQUEST
    #         )
    #     serializer = self.get_serializer(instance, data=request.data, partial=partial)
    #     if serializer.is_valid():
    #         self.perform_update(serializer)
    #         return Response(serializer.data, status=status.HTTP_201_CREATED)
    #     else:
    #         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ConfigurationsView(APIView):
    def get(self, request):
        queryset = Threshold.get_solo()
        return Response({'threshold': queryset.min_files, "show_error": queryset.show_file_error},
                        status=status.HTTP_200_OK)
